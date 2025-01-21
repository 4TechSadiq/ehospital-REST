# from docx import Document
# from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
# from docx.shared import Pt
# from fpdf import FPDF
# import os

# def generate_medical_invoice(data, output_docx='Medical_Invoice.docx', output_pdf='Medical_Invoice.pdf'):
#     """
#     Generates a medical invoice Word document and converts it to a PDF file.

#     Args:
#         data (dict): A dictionary containing medical and user information.
#         output_docx (str): The filename for the generated Word document.
#         output_pdf (str): The filename for the generated PDF file.

#     Returns:
#         None
#     """
#     # Create a new Word document
#     doc = Document()

#     # Title
#     title = doc.add_heading('Medical Invoice', level=1)
#     title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

#     # Add patient details
#     doc.add_heading('Patient Details', level=2)
#     patient_info = data['userData']
#     doc.add_paragraph(f"User ID: {patient_info['userId']}")
#     doc.add_paragraph(f"Name: {patient_info['firstName'].title()} {patient_info['lastName'].title()}")
#     doc.add_paragraph(f"Gender: {patient_info['gender'].capitalize()}")
#     doc.add_paragraph(f"Email: {patient_info['email']}")
#     doc.add_paragraph(f"Phone: {patient_info['phone']}")
#     doc.add_paragraph(f"Address: {patient_info['address']}")

#     # Add medical condition details
#     doc.add_heading('Medical Condition', level=2)
#     medical_info = data['medicalCondition']
#     doc.add_paragraph(f"Condition: {medical_info['condition'].capitalize()}")
#     doc.add_paragraph(f"Severity: {medical_info['severity']}")
#     doc.add_paragraph(f"Medication: {medical_info['medication']}")
#     doc.add_paragraph(f"Doctor Assigned: {medical_info['doctor']}")
#     doc.add_paragraph(f"Status: {medical_info['status']}")

#     # Add timestamp
#     timestamp = data['timestamp']
#     doc.add_paragraph(f"Timestamp: {timestamp}")

#     # Save the document
#     doc.save(output_docx)
#     print(f"Invoice has been generated as '{output_docx}'.")

#     # Convert the DOCX to PDF
#     pdf = FPDF()
#     pdf.set_auto_page_break(auto=True, margin=15)
#     pdf.add_page()
#     pdf.set_font("Arial", size=12)

#     # Read content from the DOCX file
#     doc = Document(output_docx)
#     for paragraph in doc.paragraphs:
#         pdf.multi_cell(0, 10, paragraph.text)

#     # Save as PDF
#     pdf.output(output_pdf)
#     print(f"PDF has been generated as '{output_pdf}'.")

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
from fpdf import FPDF
from io import BytesIO

def generate_medical_invoice(data):
    """
    Generates a medical invoice as both a Word document and a PDF in memory.

    Args:
        data (dict): A dictionary containing medical and user information.

    Returns:
        tuple: A tuple containing the Word document as a BytesIO object and the PDF as a BytesIO object.
    """
    # Create a new Word document
    doc = Document()

    # Title
    title = doc.add_heading('Medical Invoice', level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Add patient details
    doc.add_heading('Patient Details', level=2)
    patient_info = data['userData']
    doc.add_paragraph(f"User ID: {patient_info['userId']}")
    doc.add_paragraph(f"Name: {patient_info['firstName'].title()} {patient_info['lastName'].title()}")
    doc.add_paragraph(f"Gender: {patient_info['gender'].capitalize()}")
    doc.add_paragraph(f"Email: {patient_info['email']}")
    doc.add_paragraph(f"Phone: {patient_info['phone']}")
    doc.add_paragraph(f"Address: {patient_info['address']}")

    # Add medical condition details
    doc.add_heading('Medical Condition', level=2)
    medical_info = data['medicalCondition']
    doc.add_paragraph(f"Condition: {medical_info['condition'].capitalize()}")
    doc.add_paragraph(f"Severity: {medical_info['severity']}")
    doc.add_paragraph(f"Medication: {medical_info['medication']}")
    doc.add_paragraph(f"Doctor Assigned: {medical_info['doctor']}")
    doc.add_paragraph(f"Status: {medical_info['status']}")

    # Add timestamp
    timestamp = data['timestamp']
    doc.add_paragraph(f"Timestamp: {timestamp}")

    # Save the Word document to a BytesIO object
    word_buffer = BytesIO()
    doc.save(word_buffer)
    word_buffer.seek(0)

    # Create a PDF in memory
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font("Arial", size=12)

    # Read content from the Word document in memory
    doc = Document(word_buffer)
    for paragraph in doc.paragraphs:
        pdf.multi_cell(0, 10, paragraph.text)

    # Save the PDF to a BytesIO object
    pdf_buffer = BytesIO()
    pdf_content = pdf.output(dest='S').encode('latin1')  # Get the PDF content as bytes
    pdf_buffer.write(pdf_content)
    pdf_buffer.seek(0)

    return word_buffer, pdf_buffer
